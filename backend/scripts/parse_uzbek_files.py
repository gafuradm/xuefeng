# backend/scripts/parse_uzbek_files.py
import os
import re
import json
from pathlib import Path
from collections import defaultdict
import pdfplumber
from docx import Document

RAW_DIR = Path(__file__).parent.parent / "data" / "uzbek" / "raw"
OUTPUT_JSON = Path(__file__).parent.parent / "data" / "uzbek" / "uzbek_tasks.json"

# Маппинг предметов по названиям файлов
SUBJECT_MAP = {
    "matem": "математика",
    "fizika": "физика",
    "kimyo": "химия",
    "biologiya": "биология",
    "tarix": "история",
    "ona_tili": "узбекский язык",
    "rus_tili": "русский язык",
    "qqsha": "каракалпакский язык",
    "adabiyot": "литература"
}

def detect_subject(filename):
    """Определяет предмет по имени файла"""
    name = filename.lower()
    for key, subj in SUBJECT_MAP.items():
        if key in name:
            return subj
    return "общий"

def detect_language(text):
    """Определяет язык текста по характерным буквам"""
    text_lower = text[:500].lower()
    # Каракалпакские специфические буквы
    if any(c in text_lower for c in ['á', 'ı', 'ñ', 'ó', 'ú', 'ǵ', 'sh', 'ch']):
        return "kaa"  # каракалпакский
    # Узбекские специфические буквы
    elif any(c in text_lower for c in ['‘', 'ʻ', 'g‘', 'o‘']):
        return "uz"
    else:
        return "ru"

def extract_text_from_docx(filepath):
    """Извлекает текст из DOCX с поддержкой таблиц"""
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                lines.append(" | ".join(row_text))
    return "\n".join(lines)

def extract_text(filepath):
    """Извлекает текст из DOCX или PDF"""
    try:
        if filepath.suffix == ".pdf":
            with pdfplumber.open(filepath) as pdf:
                pages_text = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages_text.append(t)
                return "\n".join(pages_text)
        elif filepath.suffix == ".docx":
            return extract_text_from_docx(filepath)
    except Exception as e:
        print(f"  Ошибка извлечения: {e}")
    return ""

def clean_text(text):
    """Очищает текст от мусора"""
    # Удаляем номера страниц
    text = re.sub(r'===== Page \d+ =====', '', text)
    # Удаляем пустые строки
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    return '\n'.join(lines)

def extract_options(block, lang):
    """Извлекает варианты ответов A) B) C) D) или A. B. C. D."""
    options = {}
    
    # Паттерны для извлечения вариантов
    patterns = [
        r'([A-D])[\.\)]\s*([^\n]+)',  # A) текст или A. текст
        r'\(([A-D])\)\s*([^\n]+)',     # (A) текст
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, block)
        for letter, opt_text in matches:
            opt_text = opt_text.strip()
            if letter not in options and len(opt_text) > 1:
                options[letter] = opt_text
        if options:
            break
    
    return options

def extract_answer(block, options, lang):
    """Извлекает правильный ответ"""
    # Паттерны для поиска ответа
    patterns = [
        # Русский
        r'Ответ:\s*([A-D](?:,\s*[A-D])*)',
        r'Ответ\s+([A-D])',
        r'Правильный ответ:\s*([A-D])',
        # Узбекский
        r'Javob:\s*([A-D](?:,\s*[A-D])*)',
        r'Javob\s+([A-D])',
        # Каракалпакский
        r'Juwap:\s*([A-D](?:,\s*[A-D])*)',
        r'Juwab:\s*([A-D](?:,\s*[A-D])*)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, block, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    # Если не нашли, проверяем по вариантам (часто ответ выделен жирным или в конце)
    if options:
        # Ищем в последних 500 символах блока
        end_part = block[-500:]
        for letter in options.keys():
            if f' {letter} ' in end_part or f'\n{letter}\n' in end_part:
                return letter
    
    return ""

def extract_question_text(block, options):
    """Извлекает чистый текст вопроса без вариантов ответов"""
    q_text = block
    # Удаляем строки с ответом
    q_text = re.sub(r'(Ответ|Javob|Juwap|Juwab):\s*[A-D,\s]+', '', q_text, flags=re.IGNORECASE)
    # Удаляем варианты ответов
    for letter in options.keys():
        q_text = re.sub(rf'{letter}[\.\)]\s*[^\n]+', '', q_text)
        q_text = re.sub(rf'\({letter}\)\s*[^\n]+', '', q_text)
    # Чистим лишние пробелы
    q_text = re.sub(r'\s+', ' ', q_text).strip()
    # Удаляем слишком короткие вопросы
    if len(q_text) < 15:
        return ""
    return q_text[:2000]

def determine_topic(block, subject, lang):
    """Определяет тему задачи"""
    block_lower = block.lower()
    
    if subject == "математика":
        if any(w in block_lower for w in ['tenglama', 'уравнение', 'equation', 'x', 'y']):
            return "алгебра"
        elif any(w in block_lower for w in ['geometriya', 'геометрия', 'triangle', 'circle', 'angle']):
            return "геометрия"
        elif any(w in block_lower for w in ['trigonometriya', 'тригонометрия', 'sin', 'cos', 'tan']):
            return "тригонометрия"
        elif any(w in block_lower for w in ['hosila', 'производная', 'integral', 'limit']):
            return "математический анализ"
        elif any(w in block_lower for w in ['ehtimol', 'вероятность', 'probability']):
            return "теория вероятностей"
        else:
            return "алгебра"
    
    elif subject == "физика":
        if any(w in block_lower for w in ['tezlik', 'скорость', 'velocity', 'masofa', 'расстояние']):
            return "кинематика"
        elif any(w in block_lower for w in ['kuch', 'сила', 'force', 'mass', 'масса']):
            return "динамика"
        elif any(w in block_lower for w in ['issiqlik', 'теплота', 'temperature', 'harorat']):
            return "термодинамика"
        elif any(w in block_lower for w in ['tok', 'ток', 'current', 'kuchlanish', 'напряжение']):
            return "электричество"
        else:
            return "общая физика"
    
    elif subject == "химия":
        if any(w in block_lower for w in ['reaksiya', 'реакция', 'reaction']):
            return "химические реакции"
        elif any(w in block_lower for w in ['kislota', 'кислота', 'asos', 'основание']):
            return "кислоты и основания"
        elif any(w in block_lower for w in ['eritma', 'раствор', 'concentration']):
            return "растворы"
        else:
            return "общая химия"
    
    elif subject in ["узбекский язык", "русский язык", "каракалпакский язык"]:
        return "грамматика"
    
    elif subject in ["литература"]:
        return "анализ текста"
    
    return "общая"

def determine_difficulty(block):
    """Определяет сложность задачи"""
    if len(block) < 300:
        return "easy"
    elif len(block) > 800:
        return "hard"
    return "medium"

def split_into_questions(text):
    """Разделяет текст на отдельные вопросы по номерам"""
    # Ищем номера вопросов в начале строки
    pattern = r'(?:^|\n)(\d+)[\.\)]\s*(.*?)(?=\n\d+[\.\)]|\n\n\s*\d+[\.\)]|\Z)'
    questions = []
    matches = re.finditer(pattern, text, re.DOTALL)
    
    for match in matches:
        num = int(match.group(1))
        content = match.group(2).strip()
        if content:
            questions.append({
                "number": num,
                "content": content
            })
    return questions

def parse_file(filepath):
    """Парсит один файл и возвращает список задач"""
    print(f"  Обработка: {filepath.name}")
    
    # Извлекаем текст
    text = extract_text(filepath)
    if not text:
        print(f"    ❌ Не удалось извлечь текст")
        return []
    
    text = clean_text(text)
    subject = detect_subject(filepath.name)
    lang = detect_language(text)
    
    # Сохраняем debug текст для анализа
    debug_path = RAW_DIR.parent / f"{filepath.stem}_debug.txt"
    with open(debug_path, 'w', encoding='utf-8') as df:
        df.write(text)
    
    # Разделяем на вопросы
    questions = split_into_questions(text)
    print(f"    Найдено вопросов: {len(questions)}")
    
    tasks = []
    for q in questions:
        # Извлекаем варианты
        options = extract_options(q['content'], lang)
        # Извлекаем ответ
        answer = extract_answer(q['content'], options, lang)
        # Извлекаем текст вопроса
        question_text = extract_question_text(q['content'], options)
        
        if not question_text:
            continue
        
        # Определяем тему и сложность
        topic = determine_topic(q['content'], subject, lang)
        difficulty = determine_difficulty(q['content'])
        
        # Определяем тип вопроса
        if options:
            q_type = "multiple_correct" if (',' in answer and len(answer) > 1) else "single_correct"
        else:
            q_type = "numerical"
        
        task = {
            "number": q['number'],
            "subject": subject,
            "language": lang,
            "topic": topic,
            "difficulty": difficulty,
            "question_type": q_type,
            "question": question_text,
            "answer": answer,
            "options": options,
            "full_text": q['content'][:3000]
        }
        tasks.append(task)
    
    print(f"    Успешно обработано: {len(tasks)} задач")
    return tasks

def main():
    # Находим все файлы
    files = list(RAW_DIR.glob("*.pdf")) + list(RAW_DIR.glob("*.docx"))
    print(f"Найдено файлов: {len(files)}")
    
    all_tasks = []
    
    for f in files:
        tasks = parse_file(f)
        all_tasks.extend(tasks)
    
    # Удаляем дубликаты
    unique = {}
    for t in all_tasks:
        # Ключ: номер вопроса + предмет + первые 100 символов вопроса
        key = (t['number'], t['subject'], t['question'][:100])
        if key not in unique:
            unique[key] = t
    
    final = list(unique.values())
    final.sort(key=lambda x: (x['subject'], x['number']))
    
    # Добавляем ID
    for i, task in enumerate(final, 1):
        task["id"] = f"uzbek_{i}"
    
    # Сохраняем результат
    OUTPUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(final, f, ensure_ascii=False, indent=2)
    
    print(f"\n{'='*50}")
    print(f"✅ Сохранено задач: {len(final)}")
    
    # Статистика по предметам
    stats = defaultdict(int)
    for t in final:
        stats[t['subject']] += 1
    for subj, cnt in sorted(stats.items()):
        print(f"   {subj}: {cnt}")
    
    # Статистика по языкам
    lang_stats = defaultdict(int)
    for t in final:
        lang_stats[t['language']] += 1
    print(f"\n   По языкам:")
    for lang, cnt in lang_stats.items():
        lang_name = {"ru": "русский", "uz": "узбекский", "kaa": "каракалпакский"}.get(lang, lang)
        print(f"      {lang_name}: {cnt}")
    
    print(f"{'='*50}")

if __name__ == "__main__":
    main()