# backend/scripts/parse_india_files.py
import os
import re
import json
from pathlib import Path
import pdfplumber
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

RAW_DIR = Path(__file__).parent.parent / "data" / "india" / "raw"
OUTPUT_JSON = Path(__file__).parent.parent / "data" / "india" / "india_tasks.json"

def extract_text_from_docx(filepath):
    """Расширенное извлечение текста из DOCX с таблицами"""
    doc = Document(filepath)
    full_text = []
    
    # Извлекаем текст из параграфов
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

def extract_text(filepath):
    """Извлечение текста из DOCX или PDF"""
    try:
        if filepath.suffix == ".pdf":
            with pdfplumber.open(filepath) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            return text
        elif filepath.suffix == ".docx":
            return extract_text_from_docx(filepath)
    except Exception as e:
        print(f"Ошибка при извлечении: {e}")
    return ""

def extract_question_blocks(text):
    """Извлекает блоки вопросов, начиная с Q.1, Q.2 и т.д."""
    blocks = []
    # Ищем все Q.1, Q.2 и т.д.
    pattern = r'(Q\.\s*\d+[.\s]*.*?)(?=Q\.\s*\d+|\n\n\s*SECTION|\Z)'
    matches = re.finditer(pattern, text, re.DOTALL | re.IGNORECASE)
    
    for match in matches:
        block = match.group(1).strip()
        if block:
            blocks.append(block)
    return blocks

def parse_question_block(block, subject, filename):
    """Парсит один блок вопроса"""
    # Извлекаем номер вопроса
    num_match = re.search(r'Q\.\s*(\d+)', block, re.IGNORECASE)
    if not num_match:
        return None
    number = int(num_match.group(1))
    
    # Извлекаем текст вопроса (часть до вариантов ответов)
    # Варианты ответов обычно начинаются с A) или (A)
    options_start = re.search(r'\n\s*[A-D][\.\)]|\n\s*\([A-D]\)', block)
    if options_start:
        question_text = block[:options_start.start()].strip()
        options_part = block[options_start.start():]
    else:
        question_text = block.strip()
        options_part = ""
    
    # Извлекаем варианты ответов
    options = {}
    opt_pattern = r'([A-D])[\.\)]\s*([^\n]+)'
    for match in re.finditer(opt_pattern, options_part):
        letter = match.group(1)
        opt_text = match.group(2).strip()
        options[letter] = opt_text
    
    # Извлекаем ответ
    answer = ""
    answer_match = re.search(r'Answer\s*:\s*([A-D][,\sA-D]*)', block, re.IGNORECASE)
    if answer_match:
        answer = answer_match.group(1).strip()
    
    # Если не нашли, ищем в конце блока
    if not answer:
        end_match = re.search(r'Answer\s+([A-D])', block, re.IGNORECASE)
        if end_match:
            answer = end_match.group(1)
    
    # Очищаем вопрос от лишнего
    question_clean = re.sub(r'\s+', ' ', question_text).strip()
    
    # Определяем тип вопроса
    if len(options) > 0:
        if ',' in answer and len(answer) > 1:
            q_type = "multiple_correct"
        else:
            q_type = "single_correct"
    else:
        q_type = "numerical"
    
    # Определяем тему
    topic = determine_topic(question_clean, subject)
    difficulty = determine_difficulty(question_clean)
    
    return {
        "number": number,
        "question": question_clean[:2000],
        "full_text": block[:3000],
        "options": options,
        "answer": answer,
        "question_type": q_type,
        "topic": topic,
        "difficulty": difficulty
    }

def determine_topic(text, subject):
    text_lower = text.lower()
    if subject == "математика":
        if any(w in text_lower for w in ['matrix', 'determinant']):
            return "линейная алгебра"
        elif any(w in text_lower for w in ['limit', 'derivative', 'integral']):
            return "математический анализ"
        elif any(w in text_lower for w in ['probability', 'random']):
            return "теория вероятностей"
        elif any(w in text_lower for w in ['sin', 'cos', 'tan']):
            return "тригонометрия"
        elif any(w in text_lower for w in ['ellipse', 'parabola', 'hyperbola']):
            return "конические сечения"
        elif any(w in text_lower for w in ['area', 'volume', 'triangle']):
            return "геометрия"
        return "алгебра"
    elif subject == "физика":
        if any(w in text_lower for w in ['electric', 'charge', 'field']):
            return "электричество"
        elif any(w in text_lower for w in ['oscillation', 'spring']):
            return "колебания"
        elif any(w in text_lower for w in ['mirror', 'lens']):
            return "оптика"
        elif any(w in text_lower for w in ['kinetic', 'energy']):
            return "механика"
        return "физика"
    elif subject == "химия":
        if any(w in text_lower for w in ['organic', 'benzene']):
            return "органическая химия"
        elif any(w in text_lower for w in ['complex', 'ligand']):
            return "координационная химия"
        return "неорганическая химия"
    return "general"

def determine_difficulty(text):
    length = len(text)
    if length < 300:
        return "easy"
    elif length > 800:
        return "hard"
    return "medium"

def detect_subject_from_filename(filename):
    name = filename.lower()
    if "math" in name or "mathematics" in name:
        return "математика"
    elif "physics" in name:
        return "физика"
    elif "chemistry" in name:
        return "химия"
    return "математика"

def main():
    files = list(RAW_DIR.glob("*.docx")) + list(RAW_DIR.glob("*.pdf"))
    print(f"Найдено файлов: {len(files)}")
    all_tasks = []
    task_id = 0
    
    for f in files:
        print(f"\nОбработка: {f.name}")
        text = extract_text(f)
        
        if not text:
            print(f"  ❌ Не удалось извлечь текст")
            continue
        
        print(f"  Длина текста: {len(text)} символов")
        
        # Сохраняем текст для отладки
        debug_path = RAW_DIR.parent / f"{f.stem}_debug.txt"
        with open(debug_path, 'w', encoding='utf-8') as df:
            df.write(text)
        print(f"  📄 Сохранен debug-текст в {debug_path}")
        
        subject = detect_subject_from_filename(f.name)
        question_blocks = extract_question_blocks(text)
        print(f"  Найдено блоков вопросов: {len(question_blocks)}")
        
        for block in question_blocks:
            parsed = parse_question_block(block, subject, f.name)
            if parsed:
                task_id += 1
                task = {
                    "id": f"india_{task_id}",
                    "source": f.name,
                    "number": parsed["number"],
                    "subject": subject,
                    "topic": parsed["topic"],
                    "difficulty": parsed["difficulty"],
                    "question_type": parsed["question_type"],
                    "question": parsed["question"],
                    "answer": parsed["answer"],
                    "options": parsed["options"],
                    "solution": "",
                    "full_text": parsed["full_text"]
                }
                all_tasks.append(task)
                print(f"    Вопрос {parsed['number']}: {parsed['question_type']}")
        
        print(f"  Извлечено вопросов: {len([t for t in all_tasks if t['source'] == f.name])}")
    
    # Сохраняем результат
    OUTPUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(all_tasks, f, ensure_ascii=False, indent=2)
    
    print(f"\n{'='*50}")
    print(f"✅ Сохранено задач: {len(all_tasks)}")
    
    # Статистика по предметам
    stats = {}
    for t in all_tasks:
        subj = t['subject']
        stats[subj] = stats.get(subj, 0) + 1
    for subj, count in stats.items():
        print(f"   {subj}: {count}")
    print(f"{'='*50}")

if __name__ == "__main__":
    main()